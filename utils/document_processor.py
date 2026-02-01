"""
Document Processor for PDF and DOCX Files
Handles text extraction from uploaded documents for use in queries.
"""

import io
from typing import Optional, Dict, List
import tempfile
import os

class DocumentProcessor:
    """
    Processes uploaded documents (PDF and DOCX) to extract text content.
    Supports both single file and batch processing.
    """
    
    # Maximum file size (10 MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc'}
    
    @staticmethod
    def is_supported_file(filename: str) -> bool:
        """
        Check if file type is supported.
        
        Args:
            filename: Name of the file
            
        Returns:
            True if file type is supported, False otherwise
        """
        ext = os.path.splitext(filename.lower())[1]
        return ext in DocumentProcessor.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_content: Binary content of the PDF file
            
        Returns:
            Extracted text as string
        """
        try:
            import PyPDF2
            
            # Create a PDF reader object
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{text}")
                except Exception as e:
                    text_parts.append(f"--- Page {page_num + 1} ---\n[Error extracting text: {str(e)}]")
            
            return "\n\n".join(text_parts) if text_parts else "[No text could be extracted from PDF]"
            
        except ImportError:
            return "[Error: PyPDF2 not installed. Run: pip install PyPDF2]"
        except Exception as e:
            return f"[Error processing PDF: {str(e)}]"
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: Binary content of the DOCX file
            
        Returns:
            Extracted text as string
        """
        try:
            import docx
            
            # Create a temporary file for python-docx
            # (python-docx requires a file path, not bytes)
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(file_content)
                tmp_path = tmp_file.name
            
            try:
                # Load the document
                doc = docx.Document(tmp_path)
                
                # Extract text from paragraphs
                text_parts = []
                
                # Extract from main document body
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                
                # Extract from tables
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            table_text.append(row_text)
                    if table_text:
                        text_parts.append("\n[Table]\n" + "\n".join(table_text))
                
                return "\n\n".join(text_parts) if text_parts else "[No text found in DOCX]"
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(tmp_path)
                except:
                    pass
            
        except ImportError:
            return "[Error: python-docx not installed. Run: pip install python-docx]"
        except Exception as e:
            return f"[Error processing DOCX: {str(e)}]"
    
    @staticmethod
    def process_document(uploaded_file) -> Dict[str, str]:
        """
        Process an uploaded document and extract text.
        
        Args:
            uploaded_file: Streamlit UploadedFile object
            
        Returns:
            Dictionary with 'filename', 'content', 'error', and 'file_type'
        """
        result = {
            'filename': uploaded_file.name,
            'content': '',
            'error': None,
            'file_type': None,
            'size': uploaded_file.size
        }
        
        # Check file size
        if uploaded_file.size > DocumentProcessor.MAX_FILE_SIZE:
            result['error'] = f"File too large ({uploaded_file.size / (1024*1024):.1f} MB). Maximum size is 10 MB."
            return result
        
        # Get file extension
        ext = os.path.splitext(uploaded_file.name.lower())[1]
        result['file_type'] = ext
        
        # Check if supported
        if not DocumentProcessor.is_supported_file(uploaded_file.name):
            result['error'] = f"Unsupported file type: {ext}. Supported types: PDF, DOCX"
            return result
        
        # Read file content
        try:
            file_content = uploaded_file.read()
        except Exception as e:
            result['error'] = f"Error reading file: {str(e)}"
            return result
        
        # Extract text based on file type
        try:
            if ext == '.pdf':
                result['content'] = DocumentProcessor.extract_text_from_pdf(file_content)
            elif ext in ['.docx', '.doc']:
                result['content'] = DocumentProcessor.extract_text_from_docx(file_content)
            else:
                result['error'] = f"Unsupported file type: {ext}"
                
        except Exception as e:
            result['error'] = f"Error processing document: {str(e)}"
        
        return result
    
    @staticmethod
    def process_multiple_documents(uploaded_files: List) -> List[Dict[str, str]]:
        """
        Process multiple uploaded documents.
        
        Args:
            uploaded_files: List of Streamlit UploadedFile objects
            
        Returns:
            List of result dictionaries
        """
        results = []
        for uploaded_file in uploaded_files:
            result = DocumentProcessor.process_document(uploaded_file)
            results.append(result)
        return results
    
    @staticmethod
    def format_document_content(results: List[Dict[str, str]], 
                               include_filenames: bool = True,
                               separator: str = "\n\n" + "="*80 + "\n\n") -> str:
        """
        Format extracted document content for inclusion in queries.
        
        Args:
            results: List of processing results from process_multiple_documents
            include_filenames: Whether to include filename headers
            separator: Separator between documents
            
        Returns:
            Formatted text string ready for LLM input
        """
        formatted_parts = []
        
        for result in results:
            if result['error']:
                if include_filenames:
                    formatted_parts.append(
                        f"[File: {result['filename']}]\n"
                        f"Error: {result['error']}"
                    )
            elif result['content']:
                if include_filenames:
                    formatted_parts.append(
                        f"[Document: {result['filename']}]\n\n"
                        f"{result['content']}"
                    )
                else:
                    formatted_parts.append(result['content'])
        
        return separator.join(formatted_parts)
    
    @staticmethod
    def get_document_summary(result: Dict[str, str]) -> str:
        """
        Get a brief summary of document processing result.
        
        Args:
            result: Processing result dictionary
            
        Returns:
            Human-readable summary string
        """
        if result['error']:
            return f"❌ {result['filename']}: {result['error']}"
        
        content_length = len(result['content'])
        word_count = len(result['content'].split())
        size_mb = result['size'] / (1024 * 1024)
        
        return (f"✅ {result['filename']} "
                f"({size_mb:.2f} MB, ~{word_count:,} words, "
                f"{content_length:,} characters)")
